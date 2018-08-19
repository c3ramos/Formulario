from docx import Document
from docx.shared import Inches,Cm

doc = Document('plantillaRetilap.docx')
tables = doc.tables
p = tables[2].rows[0].cells[0].add_paragraph()
r = p.add_run()
r.add_picture('hotRide.jpg',
		width=Cm(9.94),
		height=Cm(6.07))
doc.save('addImage.docx')

