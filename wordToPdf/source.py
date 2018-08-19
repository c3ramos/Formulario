from __future__ import print_function
from mailmerge import MailMerge
from datetime import date

from docx import Document
from docx.shared import Cm

#Cargamos la plantilla
template = "plantillaRetilap.docx"
#Creamos un documento de tipo MailMerge
document = MailMerge(template)

#print(document.get_merge_fields())
#input("Presiona cualquier tecla para continuar")

#Hacemos el merge de valores

document.merge (nombreFamiliaRetilap='Familia Continua',
		codigoRevision='a12312as',
		nombreFamiliaPie='FamiliaCont')

e_wall = [{
	'itemCode': '01',
	'SubfamilyName':'FamiliaContinua',
	'color' : 'Blanco Mate'
	
	},{

	'itemCode': '02',
	'SubfamilyName':'FamiliaContinua',
	'color' : 'Negro Te Mate'
	
	}]

kubik = [{
	'itemCode2': '01',
	'SubfamilyName2':'Familasdtinua'
	},{
	'itemCode2': '02',
	'SubfamilyName2':'asdasdaaaa'	
	}]



#Las llaves en el diccionario corresponden al 
#MergeField en el doc de word

document.merge_rows('itemCode',e_wall)
document.merge_rows('itemCode2',kubik)
document.write('testOut.docx')

#Agregando imagenes a contenedores
doc = Document('testOut.docx')
tables = doc.tables
p = tables[2].rows[0].cells[0].add_paragraph()
r = p.add_run()
r.add_picture('hotRide.jpg',
		width=Cm(9.5), 
		height=Cm(5.7))

doc.save('addImage.docx')








